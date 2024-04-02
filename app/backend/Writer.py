from docx import Document
import os

def saveNotes(docxPath, text):
    if os.path.exists(docxPath):
        doc = Document(docxPath)
    else:
        doc = Document()
    doc.add_paragraph(text)
    doc.save(docxPath)
